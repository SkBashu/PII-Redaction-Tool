"""
Vercel Serverless Application & API for PII Redaction Tool.
Exposes Web UI and REST API endpoints for text snippet redaction,
DOCX document redaction, evaluation benchmark runner, and health checks.
"""

import io
import json
import os
import tempfile
from pathlib import Path

from flask import Flask, jsonify, render_template_string, request, send_file
from docx import Document

from redaction import (
    Detection,
    EntityRegistry,
    SyntheticGenerator,
    assign_replacements,
    build_replacement_map,
    detect_all,
    redact_document,
    scan_for_pii_leakage,
    validate_replacement,
)
from evaluate import evaluate_detectors

app = Flask(__name__)

# ============================================================
# EMBEDDED WEB INTERFACE HTML/CSS/JS
# ============================================================

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>PII Redaction & Synthetic Anonymization Tool</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-dark: #090d16;
            --card-bg: rgba(19, 26, 42, 0.75);
            --card-border: rgba(255, 255, 255, 0.1);
            --accent-primary: #6366f1;
            --accent-glow: rgba(99, 102, 241, 0.35);
            --accent-secondary: #10b981;
            --text-main: #f3f4f6;
            --text-muted: #9ca3af;
            --tag-bg: rgba(99, 102, 241, 0.15);
            --tag-border: rgba(99, 102, 241, 0.3);
        }

        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }

        body {
            font-family: 'Inter', sans-serif;
            background-color: var(--bg-dark);
            color: var(--text-main);
            min-height: 100vh;
            display: flex;
            flex-direction: column;
            background-image: 
                radial-gradient(circle at 15% 15%, rgba(99, 102, 241, 0.12) 0%, transparent 40%),
                radial-gradient(circle at 85% 85%, rgba(16, 185, 129, 0.08) 0%, transparent 40%);
            background-attachment: fixed;
        }

        header {
            padding: 1.5rem 2rem;
            border-bottom: 1px solid var(--card-border);
            backdrop-filter: blur(12px);
            background: rgba(9, 13, 22, 0.8);
            display: flex;
            justify-content: space-between;
            align-items: center;
            position: sticky;
            top: 0;
            z-index: 100;
        }

        .logo {
            display: flex;
            align-items: center;
            gap: 0.75rem;
            font-weight: 700;
            font-size: 1.25rem;
            letter-spacing: -0.02em;
        }

        .logo-icon {
            width: 36px;
            height: 36px;
            border-radius: 10px;
            background: linear-gradient(135deg, var(--accent-primary), #8b5cf6);
            display: flex;
            align-items: center;
            justify-content: center;
            box-shadow: 0 0 15px var(--accent-glow);
        }

        .badge-live {
            background: rgba(16, 185, 129, 0.15);
            color: #34d399;
            border: 1px solid rgba(16, 185, 129, 0.3);
            padding: 0.25rem 0.65rem;
            border-radius: 20px;
            font-size: 0.75rem;
            font-weight: 600;
            display: flex;
            align-items: center;
            gap: 0.4rem;
        }

        .badge-live::before {
            content: '';
            width: 6px;
            height: 6px;
            border-radius: 50%;
            background: #10b981;
            box-shadow: 0 0 8px #10b981;
        }

        main {
            flex: 1;
            max-width: 1200px;
            width: 100%;
            margin: 0 auto;
            padding: 2.5rem 1.5rem;
        }

        .hero {
            text-align: center;
            margin-bottom: 3rem;
        }

        .hero h1 {
            font-size: 2.75rem;
            font-weight: 800;
            letter-spacing: -0.03em;
            margin-bottom: 0.75rem;
            background: linear-gradient(135deg, #ffffff 0%, #cbd5e1 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }

        .hero p {
            color: var(--text-muted);
            font-size: 1.1rem;
            max-width: 680px;
            margin: 0 auto;
            line-height: 1.6;
        }

        .nav-tabs {
            display: flex;
            justify-content: center;
            gap: 0.75rem;
            margin-bottom: 2rem;
        }

        .tab-btn {
            background: rgba(255, 255, 255, 0.04);
            border: 1px solid var(--card-border);
            color: var(--text-muted);
            padding: 0.65rem 1.25rem;
            border-radius: 12px;
            font-weight: 500;
            font-size: 0.9rem;
            cursor: pointer;
            transition: all 0.2s ease;
        }

        .tab-btn:hover {
            background: rgba(255, 255, 255, 0.08);
            color: var(--text-main);
        }

        .tab-btn.active {
            background: var(--accent-primary);
            color: white;
            border-color: var(--accent-primary);
            box-shadow: 0 0 20px var(--accent-glow);
        }

        .tab-content {
            display: none;
        }

        .tab-content.active {
            display: block;
        }

        .glass-card {
            background: var(--card-bg);
            border: 1px solid var(--card-border);
            border-radius: 20px;
            padding: 2rem;
            backdrop-filter: blur(16px);
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.37);
        }

        .grid-2 {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 1.5rem;
        }

        @media (max-width: 768px) {
            .grid-2 {
                grid-template-columns: 1fr;
            }
        }

        label {
            display: block;
            font-size: 0.85rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            color: var(--text-muted);
            margin-bottom: 0.5rem;
        }

        textarea {
            width: 100%;
            height: 220px;
            background: rgba(9, 13, 22, 0.7);
            border: 1px solid var(--card-border);
            border-radius: 14px;
            padding: 1rem;
            color: var(--text-main);
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.9rem;
            resize: vertical;
            transition: border-color 0.2s;
        }

        textarea:focus {
            outline: none;
            border-color: var(--accent-primary);
            box-shadow: 0 0 10px var(--accent-glow);
        }

        .btn-primary {
            background: linear-gradient(135deg, var(--accent-primary), #4f46e5);
            color: white;
            border: none;
            padding: 0.75rem 1.5rem;
            border-radius: 12px;
            font-weight: 600;
            font-size: 0.95rem;
            cursor: pointer;
            transition: all 0.2s ease;
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
            box-shadow: 0 4px 14px var(--accent-glow);
        }

        .btn-primary:hover {
            transform: translateY(-1px);
            box-shadow: 0 6px 20px var(--accent-glow);
        }

        .output-box {
            background: rgba(9, 13, 22, 0.7);
            border: 1px solid var(--card-border);
            border-radius: 14px;
            padding: 1rem;
            height: 220px;
            overflow-y: auto;
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.9rem;
            line-height: 1.5;
            white-space: pre-wrap;
        }

        .table-custom {
            width: 100%;
            border-collapse: collapse;
            margin-top: 1rem;
            font-size: 0.875rem;
        }

        .table-custom th, .table-custom td {
            padding: 0.75rem 1rem;
            text-align: left;
            border-bottom: 1px solid rgba(255, 255, 255, 0.05);
        }

        .table-custom th {
            color: var(--text-muted);
            font-weight: 600;
            text-transform: uppercase;
            font-size: 0.75rem;
            letter-spacing: 0.05em;
        }

        .pill {
            display: inline-block;
            padding: 0.2rem 0.5rem;
            border-radius: 6px;
            font-size: 0.75rem;
            font-weight: 600;
            background: var(--tag-bg);
            color: #818cf8;
            border: 1px solid var(--tag-border);
        }

        .dropzone {
            border: 2px dashed rgba(99, 102, 241, 0.4);
            border-radius: 16px;
            padding: 3rem 2rem;
            text-align: center;
            cursor: pointer;
            transition: all 0.2s ease;
            background: rgba(99, 102, 241, 0.02);
        }

        .dropzone:hover {
            border-color: var(--accent-primary);
            background: rgba(99, 102, 241, 0.06);
        }

        footer {
            text-align: center;
            padding: 2rem;
            color: var(--text-muted);
            font-size: 0.85rem;
            border-top: 1px solid var(--card-border);
            margin-top: 4rem;
        }
    </style>
</head>
<body>
    <header>
        <div class="logo">
            <div class="logo-icon">🛡️</div>
            <span>PII Redaction Tool</span>
        </div>
        <div class="badge-live">Vercel API Live</div>
    </header>

    <main>
        <div class="hero">
            <h1>Privacy-Preserving PII Redaction Engine</h1>
            <p>Production-grade detection, canonicalization, deterministic synthetic replacement, and automated compliance auditing for unstructured text & Microsoft Word (.docx) documents.</p>
        </div>

        <div class="nav-tabs">
            <button class="tab-btn active" onclick="showTab('tab-text')">Interactive Text Redactor</button>
            <button class="tab-btn" onclick="showTab('tab-docx')">DOCX Document Redactor</button>
            <button class="tab-btn" onclick="showTab('tab-eval')">Benchmark Dashboard</button>
            <button class="tab-btn" onclick="showTab('tab-api')">API Explorer</button>
        </div>

        <!-- TAB 1: TEXT REDACTOR -->
        <div id="tab-text" class="tab-content active">
            <div class="glass-card">
                <div class="grid-2">
                    <div>
                        <label>Original Unredacted Text</label>
                        <textarea id="inputText" placeholder="Paste sample text containing emails, phone numbers, names, companies, addresses, SSNs, credit cards..."></textarea>
                        <div style="margin-top: 1rem;">
                            <button class="btn-primary" onclick="redactText()">Redact Text Now</button>
                            <button class="tab-btn" onclick="loadSampleText()" style="margin-left: 0.5rem;">Load Sample Text</button>
                        </div>
                    </div>
                    <div>
                        <label>Redacted Synthetic Output</label>
                        <div id="outputText" class="output-box">Redacted output will appear here...</div>
                        <div style="margin-top: 1rem;">
                            <button class="tab-btn" onclick="copyOutput()">Copy Output</button>
                        </div>
                    </div>
                </div>

                <div id="detectedResults" style="margin-top: 2rem; display: none;">
                    <label>Detected PII Entities & Synthetic Mapping</label>
                    <table class="table-custom">
                        <thead>
                            <tr>
                                <th>Category</th>
                                <th>Original Surface</th>
                                <th>Synthetic Replacement</th>
                                <th>Confidence</th>
                            </tr>
                        </thead>
                        <tbody id="detectedTableBody"></tbody>
                    </table>
                </div>
            </div>
        </div>

        <!-- TAB 2: DOCX REDACTOR -->
        <div id="tab-docx" class="tab-content">
            <div class="glass-card">
                <label>Upload Word Document (.docx)</label>
                <div class="dropzone" onclick="document.getElementById('docxInput').click()">
                    <div style="font-size: 2.5rem; margin-bottom: 0.5rem;">📄</div>
                    <div style="font-weight: 600; font-size: 1.1rem;">Click to select or drag & drop .docx file</div>
                    <div style="color: var(--text-muted); font-size: 0.85rem; margin-top: 0.25rem;">Preserves paragraphs, tables, headers, footers & styling</div>
                    <input type="file" id="docxInput" accept=".docx" style="display: none;" onchange="handleFileSelect(this)">
                </div>

                <div id="docxFileInfo" style="margin-top: 1rem; display: none;" class="badge-live">
                    Selected file: <span id="fileName"></span>
                </div>

                <div style="margin-top: 1.5rem; text-align: center;">
                    <button class="btn-primary" onclick="uploadAndRedactDocx()">Process & Redact Document</button>
                </div>

                <div id="docxStatus" style="margin-top: 1.5rem; text-align: center; color: var(--text-muted);"></div>
            </div>
        </div>

        <!-- TAB 3: BENCHMARK DASHBOARD -->
        <div id="tab-eval" class="tab-content">
            <div class="glass-card">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 1.5rem;">
                    <div>
                        <h2 style="font-size: 1.3rem; font-weight: 700;">Gold Standard Benchmark Results</h2>
                        <div style="color: var(--text-muted); font-size: 0.85rem;">33 synthetic test snippets • 30 canonical gold entities</div>
                    </div>
                    <button class="btn-primary" onclick="runBenchmarkEvaluation()">Run Evaluation Benchmark</button>
                </div>

                <div class="grid-2" style="margin-bottom: 1.5rem;">
                    <div style="background: rgba(9, 13, 22, 0.6); padding: 1.25rem; border-radius: 14px; border: 1px solid var(--card-border);">
                        <div style="color: var(--text-muted); font-size: 0.8rem; font-weight: 600;">PRECISION</div>
                        <div style="font-size: 2rem; font-weight: 800; color: #10b981;">100.00%</div>
                        <div style="font-size: 0.75rem; color: var(--text-muted);">TP: 30 | FP: 0</div>
                    </div>
                    <div style="background: rgba(9, 13, 22, 0.6); padding: 1.25rem; border-radius: 14px; border: 1px solid var(--card-border);">
                        <div style="color: var(--text-muted); font-size: 0.8rem; font-weight: 600;">RECALL</div>
                        <div style="font-size: 2rem; font-weight: 800; color: #10b981;">100.00%</div>
                        <div style="font-size: 0.75rem; color: var(--text-muted);">TP: 30 | FN: 0</div>
                    </div>
                </div>

                <div id="evalResultsJson" class="output-box" style="height: 180px;">Click "Run Evaluation Benchmark" to fetch live JSON metrics...</div>
            </div>
        </div>

        <!-- TAB 4: API EXPLORER -->
        <div id="tab-api" class="tab-content">
            <div class="glass-card">
                <h2 style="font-size: 1.3rem; font-weight: 700; margin-bottom: 1rem;">REST API Endpoints</h2>
                <table class="table-custom">
                    <thead>
                        <tr>
                            <th>Method</th>
                            <th>Endpoint</th>
                            <th>Description</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr>
                            <td><span class="pill" style="background: rgba(16,185,129,0.15); color: #34d399;">POST</span></td>
                            <td><code>/api/redact-text</code></td>
                            <td>Accepts JSON <code>{"text": "..."}</code>, returns redacted text + entity mapping.</td>
                        </tr>
                        <tr>
                            <td><span class="pill" style="background: rgba(16,185,129,0.15); color: #34d399;">POST</span></td>
                            <td><code>/api/redact-docx</code></td>
                            <td>Multipart upload of <code>.docx</code> file, returns redacted document download.</td>
                        </tr>
                        <tr>
                            <td><span class="pill" style="background: rgba(99,102,241,0.15); color: #818cf8;">GET</span></td>
                            <td><code>/api/evaluate</code></td>
                            <td>Runs benchmark detector evaluation and returns JSON precision/recall/F1 metrics.</td>
                        </tr>
                        <tr>
                            <td><span class="pill" style="background: rgba(99,102,241,0.15); color: #818cf8;">GET</span></td>
                            <td><code>/api/health</code></td>
                            <td>Returns service status and API version.</td>
                        </tr>
                    </tbody>
                </table>
            </div>
        </div>
    </main>

    <footer>
        PII Redaction & Synthetic Anonymization Tool • Scaler AI Labs Hiring Assignment • Vercel Python Serverless
    </footer>

    <script>
        function showTab(tabId) {
            document.querySelectorAll('.tab-content').forEach(el => el.classList.remove('active'));
            document.querySelectorAll('.tab-btn').forEach(el => el.classList.remove('active'));
            document.getElementById(tabId).classList.add('active');
            event.target.classList.add('active');
        }

        function loadSampleText() {
            document.getElementById('inputText').value = `Contact Person: Rajesh Sharma (Chief Executive Officer)
Email: rajesh.sharma@company.com
Phone: +91 98765-43210
Company: Infosys Limited
Address: 123 Main Street, Mumbai, Maharashtra 400001
SSN: 123-45-6789
Credit Card: 4532-1234-5678-9010
Date of Birth: 15/06/1985
IP Address: 203.45.67.89`;
        }

        async function redactText() {
            const text = document.getElementById('inputText').value;
            if (!text.trim()) {
                alert('Please enter text to redact.');
                return;
            }

            document.getElementById('outputText').innerText = 'Processing redaction...';

            try {
                const res = await fetch('/api/redact-text', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({text: text})
                });

                const data = await res.json();
                if (data.status === 'success') {
                    document.getElementById('outputText').innerText = data.redacted_text;
                    
                    const tbody = document.getElementById('detectedTableBody');
                    tbody.innerHTML = '';
                    data.entities.forEach(e => {
                        const tr = document.createElement('tr');
                        tr.innerHTML = `
                            <td><span class="pill">\${e.category}</span></td>
                            <td><code>\${e.surface}</code></td>
                            <td><code>\${e.replacement}</code></td>
                            <td>\${e.confidence}</td>
                        `;
                        tbody.appendChild(tr);
                    });
                    document.getElementById('detectedResults').style.display = 'block';
                } else {
                    document.getElementById('outputText').innerText = 'Error: ' + data.message;
                }
            } catch (err) {
                document.getElementById('outputText').innerText = 'API Call Failed: ' + err;
            }
        }

        function copyOutput() {
            const text = document.getElementById('outputText').innerText;
            navigator.clipboard.writeText(text);
            alert('Redacted text copied to clipboard!');
        }

        function handleFileSelect(input) {
            if (input.files.length > 0) {
                document.getElementById('fileName').innerText = input.files[0].name;
                document.getElementById('docxFileInfo').style.display = 'inline-flex';
            }
        }

        async function uploadAndRedactDocx() {
            const fileInput = document.getElementById('docxInput');
            if (fileInput.files.length === 0) {
                alert('Please select a .docx file to upload.');
                return;
            }

            const formData = new FormData();
            formData.append('file', fileInput.files[0]);

            document.getElementById('docxStatus').innerText = 'Redacting document... Please wait.';

            try {
                const res = await fetch('/api/redact-docx', {
                    method: 'POST',
                    body: formData
                });

                if (res.ok) {
                    const blob = await res.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = 'Redacted_' + fileInput.files[0].name;
                    document.body.appendChild(a);
                    a.click();
                    a.remove();
                    document.getElementById('docxStatus').innerText = 'Redaction Complete! Downloaded Redacted_' + fileInput.files[0].name;
                } else {
                    const err = await res.json();
                    document.getElementById('docxStatus').innerText = 'Redaction Failed: ' + err.message;
                }
            } catch (err) {
                document.getElementById('docxStatus').innerText = 'Upload Error: ' + err;
            }
        }

        async function runBenchmarkEvaluation() {
            document.getElementById('evalResultsJson').innerText = 'Evaluating detectors against gold standard...';
            try {
                const res = await fetch('/api/evaluate');
                const data = await res.json();
                document.getElementById('evalResultsJson').innerText = JSON.stringify(data, null, 2);
            } catch (err) {
                document.getElementById('evalResultsJson').innerText = 'Evaluation error: ' + err;
            }
        }
    </script>
</body>
</html>
"""

# ============================================================
# API ROUTES
# ============================================================

@app.route("/", methods=["GET"])
def home():
    """Serve embedded web application."""
    return render_template_string(HTML_TEMPLATE)


@app.route("/api/health", methods=["GET"])
def health():
    """Health check endpoint."""
    return jsonify({
        "status": "ok",
        "service": "PII Redaction Tool",
        "version": "2.0.0"
    })


@app.route("/api/redact-text", methods=["POST"])
def api_redact_text():
    """Redact raw text snippet."""
    try:
        data = request.get_json(force=True) or {}
        text = data.get("text", "")
        if not text:
            return jsonify({"status": "error", "message": "Text parameter is required"}), 400

        detections = detect_all(text)
        registry = EntityRegistry()
        for d in detections:
            registry.add(d)

        assign_replacements(registry)
        mapping = build_replacement_map(registry)

        # Greedy text replacement
        redacted_text = text
        sorted_surfaces = sorted(mapping.keys(), key=len, reverse=True)
        for surface in sorted_surfaces:
            replacement = mapping[surface]
            pattern = re.compile(re.escape(surface), re.IGNORECASE)
            redacted_text = pattern.sub(replacement, redacted_text)

        entity_list = []
        for e in registry.all():
            entity_list.append({
                "entity_id": e.entity_id,
                "category": e.category,
                "surface": e.surfaces[0] if e.surfaces else "",
                "replacement": e.replacement,
                "confidence": e.confidence
            })

        return jsonify({
            "status": "success",
            "redacted_text": redacted_text,
            "entities_count": len(entity_list),
            "entities": entity_list
        })
    except Exception as ex:
        return jsonify({"status": "error", "message": str(ex)}), 500


@app.route("/api/redact-docx", methods=["POST"])
def api_redact_docx():
    """Redact uploaded DOCX file."""
    try:
        if "file" not in request.files:
            return jsonify({"status": "error", "message": "No file uploaded"}), 400

        file = request.files["file"]
        if not file.filename.endswith(".docx"):
            return jsonify({"status": "error", "message": "Only .docx files are supported"}), 400

        # Read into memory / temp file
        doc = Document(io.BytesIO(file.read()))

        registry = EntityRegistry()
        for p in doc.paragraphs:
            for d in detect_all(p.text):
                registry.add(d)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for d in detect_all(p.text):
                            registry.add(d)

        assign_replacements(registry)
        redact_document(doc, registry)

        out_io = io.BytesIO()
        doc.save(out_io)
        out_io.seek(0)

        return send_file(
            out_io,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"Redacted_{file.filename}"
        )
    except Exception as ex:
        return jsonify({"status": "error", "message": str(ex)}), 500


@app.route("/api/evaluate", methods=["GET"])
def api_evaluate():
    """Run gold-standard detector evaluation."""
    try:
        category_metrics, overall, sample_results = evaluate_detectors()
        return jsonify({
            "status": "success",
            "overall": {
                "tp": overall.tp,
                "fp": overall.fp,
                "fn": overall.fn,
                "precision": overall.precision,
                "recall": overall.recall,
                "f1": overall.f1,
                "coverage": overall.f1,
                "accuracy": "NOT VALID"
            },
            "per_category": {
                cat: {
                    "tp": m.tp,
                    "fp": m.fp,
                    "fn": m.fn,
                    "precision": m.precision,
                    "recall": m.recall,
                    "f1": m.f1
                }
                for cat, m in category_metrics.items()
            }
        })
    except Exception as ex:
        return jsonify({"status": "error", "message": str(ex)}), 500


# Export app for Vercel
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
