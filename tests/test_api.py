"""Synthetic-only integration tests for the Flask deployment adapter."""

from __future__ import annotations

import io
import sys
import unittest
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from app import app


class TestDeploymentAPI(unittest.TestCase):
    def setUp(self) -> None:
        self.client = app.test_client()

    def test_public_get_routes_are_safe_and_available(self) -> None:
        self.assertEqual(self.client.get("/").status_code, 200)
        self.assertEqual(self.client.get("/api/health").get_json(), {"status": "ok"})

        info = self.client.get("/api/info")
        self.assertEqual(info.status_code, 200)
        self.assertEqual(info.get_json()["service"], "PII Redaction Tool")

        demo = self.client.get("/api/demo")
        self.assertEqual(demo.status_code, 200)
        self.assertIn("riya.sharma@example.com", demo.get_json()["sample_text"])

    def test_synthetic_text_is_redacted(self) -> None:
        source = "Contact Person: Riya Sharma; email riya.sharma@example.com"
        response = self.client.post("/api/redact-text", json={"text": source})
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["status"], "success")
        self.assertNotIn("riya.sharma@example.com", payload["redacted_text"])
        self.assertGreater(payload["entities_count"], 0)

    def test_synthetic_docx_is_redacted_and_returned(self) -> None:
        document = Document()
        document.add_paragraph("Email: riya.sharma@example.com")
        document.add_paragraph("Phone: +91 9000000000")
        source = io.BytesIO()
        document.save(source)
        source.seek(0)

        response = self.client.post(
            "/api/redact-docx",
            data={"file": (source, "synthetic.docx")},
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.mimetype,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        result = Document(io.BytesIO(response.data))
        text = "\n".join(paragraph.text for paragraph in result.paragraphs)
        self.assertNotIn("riya.sharma@example.com", text)
        self.assertNotIn("+91 9000000000", text)


if __name__ == "__main__":
    unittest.main(verbosity=2)
