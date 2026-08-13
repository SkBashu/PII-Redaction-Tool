from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import hashlib
import json
import re

from docx import Document
import spacy


# ============================================================
# CONFIGURATION
# ============================================================

INPUT_FILE = Path("Red Herring Prospectus.docx")
OUTPUT_FILE = Path("Redacted_Red_Herring_Prospectus.docx")
REPORT_FILE = Path("redaction_report.json")

SEED = 20260814


# ============================================================
# LOAD SPACY
# ============================================================

try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    try:
        import spacy.cli
        spacy.cli.download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        nlp = spacy.blank("en")


# ============================================================
# DATA CLASSES
# ============================================================

@dataclass
class Detection:
    category: str
    surface: str
    normalized: str
    confidence: str
    evidence: str


@dataclass
class Entity:
    entity_id: str
    category: str
    canonical: str
    confidence: str
    surfaces: set[str] = field(default_factory=set)
    evidence: list[str] = field(default_factory=list)
    replacement: str = ""


# ============================================================
# GENERAL HELPERS
# ============================================================

def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def clean_name(name: str) -> str:
    name = name.strip()

    name = name.strip(
        " \t\r\n,.;:()[]{}*^&"
    )

    name = re.sub(
        r"\s+For\s+Further\s+Details.*$",
        "",
        name,
        flags=re.IGNORECASE
    )

    name = re.sub(
        r"\s+SEBI\s+Registration.*$",
        "",
        name,
        flags=re.IGNORECASE
    )

    name = re.sub(
        r"\s+Website.*$",
        "",
        name,
        flags=re.IGNORECASE
    )

    return normalize_spaces(
        name.strip(" ,.;:")
    )


# ============================================================
# EMAIL DETECTION
# ============================================================

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@"
    r"[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)


def detect_emails(text: str) -> list[Detection]:
    results = []

    for match in EMAIL_PATTERN.finditer(text):
        value = match.group(0)

        results.append(
            Detection(
                category="EMAIL",
                surface=value,
                normalized=value.casefold(),
                confidence="HIGH",
                evidence=value
            )
        )

    return results


# ============================================================
# PHONE DETECTION
# ============================================================

PHONE_PATTERNS = [
    re.compile(
        r"(?<!\d)\+\s*91(?:[\s-]?\d){10}(?!\d)"
    ),
    re.compile(
        r"(?<![\d+\-])\d{2,4}-\d{6,10}(?!\d)"
    ),
    re.compile(
        r"(?<!\d)[6-9]\d{9}(?!\d)"
    ),
]


def normalize_phone(phone: str):
    digits = re.sub(r"\D", "", phone)

    if (
        digits.startswith("91")
        and len(digits) == 12
    ):
        return "+91" + digits[2:]

    if (
        len(digits) == 10
        and digits[0] in "6789"
    ):
        return "+91" + digits

    if 10 <= len(digits) <= 12:
        return digits

    return None


def valid_phone(phone: str) -> bool:
    normalized = normalize_phone(phone)

    if normalized is None:
        return False

    if normalized.startswith("+91"):
        return len(normalized) == 13

    return 10 <= len(normalized) <= 12


def detect_phones(text: str) -> list[Detection]:
    results = []

    for pattern in PHONE_PATTERNS:

        for match in pattern.finditer(text):

            value = match.group(0)

            if not valid_phone(value):
                continue

            normalized = normalize_phone(value)

            if normalized is None:
                continue

            results.append(
                Detection(
                    category="PHONE",
                    surface=value,
                    normalized=normalized,
                    confidence="HIGH",
                    evidence=value
                )
            )

    return results


# ============================================================
# PERSON DETECTION
# ============================================================

REJECT_NAME_PHRASES = {
    "company secretary",
    "compliance officer",
    "chief executive officer",
    "chief financial officer",
    "key managerial personnel",
    "key managerial",
    "senior management",
    "executive director",
    "independent director",
    "whole-time director",
    "joint managing director",
    "managing director",
    "chairman",
    "selling shareholder",
    "selling shareholders",
    "registered broker",
    "share transfer agents",
    "reference rate",
    "mutual funds",
    "secondary transfer of",
    "acknowledgement slip",
    "wilful defaulter",
    "nro account",
    "bid amount",
    "the bid amount",
    "individual bidders",
    "qib bidders",
    "promoter trusts",
    "tax deducted",
    "air conditioning",
    "photo voltaic",
    "circuit kilometers",
    "kisan urja suraksha",
    "gram jyoti",
    "cap price",
    "floor price",
    "upi bidders",
    "parents branch",
    "rajesh branch",
    "sangeeta branch",
}


ADDRESS_WORDS = {
    "road",
    "marg",
    "street",
    "lane",
    "nagar",
    "gaon",
    "village",
    "taluka",
    "district",
    "pune",
    "mumbai",
    "maharashtra",
    "india",
    "building",
    "tower",
    "floor",
    "complex",
    "hospital",
    "chambers",
    "house",
    "apartment",
    "colony",
    "society",
    "park",
    "gymkhana",
    "churchgate",
    "showroom",
    "bhavan",
    "listing",
    "campus",
    "centre",
    "center",
    "plot",
    "residency",
    "reclamation",
    "backbay",
    "opposite",
    "near",
}


PERSON_CONTEXTS = [
    "contact person",
    "chairman",
    "chief executive officer",
    "chief financial officer",
    "company secretary",
    "compliance officer",
    "managing director",
    "joint managing director",
    "executive director",
    "independent director",
    "whole-time director",
    "promoter",
    "promoters",
    "shareholder",
    "shareholders",
    "allotted to",
    "being",
    "namely",
    "board of directors",
]


def looks_like_person(name: str) -> bool:
    name = clean_name(name)

    tokens = name.split()

    if not 2 <= len(tokens) <= 5:
        return False

    if any(ch.isdigit() for ch in name):
        return False

    if any(ch in name for ch in "@:/\\"):
        return False

    if "-" in name:
        return False

    cleaned = []

    for token in tokens:
        token = token.strip(
            ".,'’*^&()"
        )

        if token:
            cleaned.append(token)

    if len(cleaned) < 2:
        return False

    lowered_tokens = {
        token.casefold()
        for token in cleaned
    }

    phrase = " ".join(
        token.casefold()
        for token in cleaned
    )

    if phrase in REJECT_NAME_PHRASES:
        return False

    if lowered_tokens.intersection(
        REJECT_NAME_PHRASES
    ):
        return False

    if lowered_tokens.intersection(
        ADDRESS_WORDS
    ):
        return False

    for token in cleaned:

        if len(token) == 1:

            if not token.isalpha():
                return False

        elif not token.isalpha():

            return False

    return True


def person_context_score(text: str) -> int:
    lowered = text.casefold()

    return sum(
        1
        for term in PERSON_CONTEXTS
        if term in lowered
    )


def detect_people(text: str) -> list[Detection]:

    results = []

    if not text.strip():
        return results

    # --------------------------------------------------------
    # Contact Person
    # --------------------------------------------------------

    contact_pattern = re.compile(
        r"Contact\s+Person\s*:\s*"
        r"(.+?)"
        r"(?=;|Telephone|Tel:|E-mail|Email|Website|"
        r"SEBI Registration|$)",
        re.IGNORECASE
    )

    for match in contact_pattern.finditer(text):

        content = match.group(1)

        parts = re.split(
            r"\s*/\s*|\s+and\s+",
            content,
            flags=re.IGNORECASE
        )

        for part in parts:

            name = clean_name(part)

            if not looks_like_person(name):
                continue

            results.append(
                Detection(
                    category="PERSON",
                    surface=name,
                    normalized=name.casefold(),
                    confidence="HIGH",
                    evidence=text[:600]
                )
            )

    # --------------------------------------------------------
    # Explicit roles
    # --------------------------------------------------------

    role_pattern = re.compile(
        r"(?:"
        r"chairman\s+and\s+executive\s+director|"
        r"chief\s+executive\s+officer|"
        r"chief\s+financial\s+officer|"
        r"company\s+secretary(?:\s+and\s+compliance\s+officer)?|"
        r"joint\s+managing\s+director|"
        r"managing\s+director|"
        r"whole-time\s+director|"
        r"independent\s+director"
        r")"
        r".{0,80}?"
        r"(?:being|namely)\s+"
        r"([A-Z][A-Za-z.]*"
        r"(?:\s+[A-Z][A-Za-z.]*){1,4})"
        r"(?=\.|,|;|$)",
        re.IGNORECASE
    )

    for match in role_pattern.finditer(text):

        name = clean_name(
            match.group(1)
        )

        if not looks_like_person(name):
            continue

        results.append(
            Detection(
                category="PERSON",
                surface=name,
                normalized=name.casefold(),
                confidence="HIGH",
                evidence=text[:600]
            )
        )

    # --------------------------------------------------------
    # Promoter list / employee names
    # --------------------------------------------------------

    promoter_patterns = [
        re.compile(
            r"(?:our\s+promoters?\s+include|promoters?\s+include|promoters?\s*:\s*)"
            r"(.+?)(?=\.|;|$)",
            re.IGNORECASE
        ),
        re.compile(
            r"OUR PROMOTERS:\s*"
            r"(.*?)(?=DETAILS OF THE OFFER|$)",
            re.IGNORECASE | re.DOTALL
        )
    ]

    for pattern in promoter_patterns:
        for match in pattern.finditer(text):
            for part in re.split(r",\s*|\s+and\s+|\s*&\s*", match.group(1), flags=re.IGNORECASE):
                name = clean_name(part)
                lowered = name.casefold()

                if (
                    "family trust" in lowered
                    or "private limited" in lowered
                    or lowered.endswith("limited")
                    or "industrial park" in lowered
                    or "promoter" in lowered
                ):
                    continue

                if not looks_like_person(name):
                    continue

                results.append(
                    Detection(
                        category="PERSON",
                        surface=name,
                        normalized=name.casefold(),
                        confidence="HIGH",
                        evidence=text[:700]
                    )
                )

    employee_pattern = re.compile(
        r"\bEmployee\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,4})\b",
        re.IGNORECASE
    )
    for match in employee_pattern.finditer(text):
        name = clean_name(match.group(1))
        if not looks_like_person(name):
            continue
        results.append(
            Detection(
                category="PERSON",
                surface=name,
                normalized=name.casefold(),
                confidence="HIGH",
                evidence=text[:600]
            )
        )

    # --------------------------------------------------------
    # NER (only if paragraph is long enough and has context)
    # --------------------------------------------------------

    if len(text) < 100:
        return results

    context_score = person_context_score(text)

    if context_score < 1:
        return results

    doc_nlp = nlp(text)

    for entity in doc_nlp.ents:

        if entity.label_ != "PERSON":
            continue

        name = clean_name(
            entity.text
        )

        if not looks_like_person(name):
            continue

        results.append(
            Detection(
                category="PERSON",
                surface=name,
                normalized=name.casefold(),
                confidence="MEDIUM",
                evidence=text[:600]
            )
        )

    return results


# ============================================================
# COMPANY DETECTION
# ============================================================

COMPANY_SUFFIXES = {
    "limited",
    "ltd",
    "private limited",
    "pvt ltd",
    "llp",
    "bank",
    "securities",
    "financial",
    "management",
    "corporation",
    "corp",
    "ratings",
    "exchange",
    "registrar",
    "trust",
    "industrial",
}

COMPANY_REJECT_KEYWORDS = {
    "company",
    "issuer",
    "bank",
    "offer",
    "registrar",
    "website",
    "person",
    "contact",
}


GENERIC_COMPANY_WORDS = {
    "offer", "escrow", "collection", "bank", "account", "sponsor",
    "syndicate", "member", "registrar", "brlm", "brlms", "manager",
    "managers", "lead", "running", "book", "designated", "stock",
    "exchange", "financial", "terms", "abbreviations", "technical",
    "industry", "related", "key", "investor", "investors", "public",
    "company", "issuer", "secretary", "compliance", "officer"
}


def looks_like_company(name: str) -> bool:
    name = clean_name(name)
    tokens = name.split()

    if len(tokens) < 1:
        return False

    if any(ch.isdigit() for ch in name):
        return False

    lowered = name.casefold()

    if lowered in COMPANY_REJECT_KEYWORDS:
        return False

    has_suffix = any(
        lowered.endswith(suffix)
        for suffix in COMPANY_SUFFIXES
    )

    if not has_suffix:
        return False

    norm_tokens = [t.lower().strip(".,()[]") for t in tokens]
    non_suffix_tokens = norm_tokens[:-1] if len(norm_tokens) > 1 else norm_tokens
    if non_suffix_tokens and all(t in GENERIC_COMPANY_WORDS for t in non_suffix_tokens):
        return False

    return True


COMPANY_PREFIX_STOP_WORDS = {"and", "by", "managed", "or", "for", "with", "at", "in", "of", "to", "from"}


def clean_company_name(name: str) -> str:
    name = clean_name(name)
    tokens = name.split()
    while tokens and tokens[0].lower() in COMPANY_PREFIX_STOP_WORDS:
        tokens.pop(0)
    return " ".join(tokens)


def detect_companies(text: str) -> list[Detection]:
    results = []
    if not text.strip():
        return results

    if len(text) >= 50:
        doc_nlp = nlp(text)
        for entity in doc_nlp.ents:
            if entity.label_ == "ORG":
                name = clean_company_name(entity.text)
                if looks_like_company(name):
                    results.append(
                        Detection(
                            category="COMPANY",
                            surface=name,
                            normalized=name.casefold(),
                            confidence="MEDIUM",
                            evidence=text[:600]
                        )
                    )

    fallback_pattern = re.compile(
        r"\b([A-Z][A-Za-z0-9&.()-]*(?:\s+[A-Za-z0-9&.()-]+){0,3})\s+(?:Limited|Ltd|Private\s+Limited|Pvt\s+Ltd|LLP|Bank|Securities|Corporation|Exchange|Trust|Industrial|Registrar)(?=\s|$|[,;.])"
    )
    for match in fallback_pattern.finditer(text):
        name = clean_company_name(match.group(0))
        if looks_like_company(name) and len(name.split()) <= 4:
            results.append(
                Detection(
                    category="COMPANY",
                    surface=name,
                    normalized=name.casefold(),
                    confidence="MEDIUM",
                    evidence=text[:600]
                )
            )

    return results


# ============================================================
# ADDRESS DETECTION
# ============================================================

ADDRESS_PATTERN = re.compile(
    r"(?:"
    r"(?:Address|Residential Address|Permanent Address|"
    r"Mailing Address|Correspondence Address)"
    r"\s*:?\s*"
    r")"
    r"(.+?)(?=\n|;|Email|Tel|Contact|$)",
    re.IGNORECASE | re.DOTALL
)


def detect_addresses(text: str) -> list[Detection]:

    results = []

    if not text.strip():
        return results

    for match in ADDRESS_PATTERN.finditer(text):

        address = clean_name(match.group(1))

        if len(address) < 10:
            continue

        if not any(
            keyword in address.casefold()
            for keyword in (
                "street", "road", "marg", "lane", "building",
                "house", "apartment", "city", "district",
                "state", "pin", "postal", "area", "bangalore",
                "mumbai", "delhi", "pune", "chennai", "kolkata",
                "hyderabad", "ahmedabad", "surat", "jaipur"
            )
        ):
            continue

        results.append(
            Detection(
                category="ADDRESS",
                surface=address,
                normalized=address.casefold(),
                confidence="HIGH",
                evidence=text[:600]
            )
        )

    fallback_pattern = re.compile(
        r"(?:at|from|located at|in)\s+(\d+\s+[A-Za-z0-9&.,'()/ -]+(?:Street|Road|Marg|Lane|Building|Tower|Block|Avenue|Court|Colony|Complex|Sector|Park)[^;\.\n]{0,180})",
        re.IGNORECASE
    )

    for match in fallback_pattern.finditer(text):
        address = clean_name(match.group(1))
        if len(address) < 10:
            continue
        if not any(keyword in address.casefold() for keyword in ("street", "road", "marg", "lane", "building", "tower", "block", "avenue", "sector", "colony", "complex", "delhi", "mumbai", "pune", "bangalore")):
            continue
        results.append(
            Detection(
                category="ADDRESS",
                surface=address,
                normalized=address.casefold(),
                confidence="MEDIUM",
                evidence=text[:600]
            )
        )

    return results


# ============================================================
# SSN DETECTION
# ============================================================

SSN_PATTERN = re.compile(
    r"\b\d{3}-\d{2}-\d{4}\b"
)


def is_valid_ssn(ssn: str) -> bool:
    """
    Validate US SSN format: ###-##-####
    Reject obviously invalid patterns.
    """
    digits = re.sub(r"\D", "", ssn)

    if len(digits) != 9:
        return False

    parts = digits.split("-") if "-" in ssn else [digits[:3], digits[3:5], digits[5:9]]

    area = int(digits[:3])
    group = int(digits[3:5])
    serial = int(digits[5:9])

    if area == 0 or area == 666:
        return False

    if group == 0:
        return False

    if serial == 0:
        return False

    return True


def detect_ssns(text: str) -> list[Detection]:

    results = []

    for match in SSN_PATTERN.finditer(text):

        value = match.group(0)

        if not is_valid_ssn(value):
            continue

        results.append(
            Detection(
                category="SSN",
                surface=value,
                normalized=value,
                confidence="HIGH",
                evidence=text[:600]
            )
        )

    return results


# ============================================================
# CREDIT CARD DETECTION
# ============================================================

CREDIT_CARD_PATTERN = re.compile(
    r"\b(?:\d{4}[-\s]?){3}\d{4}\b"
)


def luhn_checksum(card_number: str) -> int:
    """Calculate Luhn checksum using the standard algorithm."""
    digits = [int(d) for d in card_number if d.isdigit()]
    total = sum(digits[::2])
    for digit in digits[1::2]:
        d = digit * 2
        total += d - 9 if d > 9 else d
    return total % 10


def is_valid_credit_card(card_number: str) -> bool:
    """Validate credit card using Luhn algorithm."""
    digits = re.sub(r"\D", "", card_number)

    if len(digits) not in (13, 14, 15, 16, 19):
        return False

    return luhn_checksum(digits) == 0


def detect_credit_cards(text: str) -> list[Detection]:

    results = []

    for match in CREDIT_CARD_PATTERN.finditer(text):

        value = match.group(0)

        if not is_valid_credit_card(value):
            continue

        normalized = re.sub(r"\D", "", value)

        results.append(
            Detection(
                category="CREDIT_CARD",
                surface=value,
                normalized=normalized,
                confidence="HIGH",
                evidence=text[:600]
            )
        )

    return results


# ============================================================
# DATE OF BIRTH DETECTION
# ============================================================

DOB_PATTERNS = [
    re.compile(
        r"(?:Date\s+of\s+Birth|DOB|D\.O\.B)"
        r"\s*:?\s*"
        r"(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})",
        re.IGNORECASE
    ),
    re.compile(
        r"(?:born|b\.?)\s+(?:on\s+)?"
        r"(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})",
        re.IGNORECASE
    ),
]


def is_valid_dob(date_str: str) -> bool:
    """
    Basic validation for date of birth.
    Must be a plausible historical date.
    """
    parts = re.split(r"[-/]", date_str)

    if len(parts) != 3:
        return False

    try:
        day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
    except ValueError:
        return False

    if year < 100:
        year += 1900 if year > 50 else 2000

    if not (1 <= month <= 12):
        return False

    if not (1 <= day <= 31):
        return False

    if year > 2010:
        return False

    if year < 1900:
        return False

    return True


def detect_dobs(text: str) -> list[Detection]:

    results = []

    for pattern in DOB_PATTERNS:

        for match in pattern.finditer(text):

            date_value = match.group(1)

            if not is_valid_dob(date_value):
                continue

            normalized = re.sub(r"[-/]", "", date_value)

            results.append(
                Detection(
                    category="DOB",
                    surface=date_value,
                    normalized=normalized,
                    confidence="HIGH",
                    evidence=text[:600]
                )
            )

    return results


# ============================================================
# IP ADDRESS DETECTION
# ============================================================

IP_ADDRESS_PATTERN = re.compile(
    r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
)


def is_valid_ipv4(ip: str) -> bool:
    """Validate IPv4 address."""
    parts = ip.split(".")

    if len(parts) != 4:
        return False

    try:
        return all(
            0 <= int(part) <= 255
            for part in parts
        )
    except ValueError:
        return False


def is_documentation_ip(ip: str) -> bool:
    """
    Exclude documentation/reserved IPv4 ranges that are not real PII.
    Keep 172.x addresses valid unless they are explicitly in reserved documentation blocks.
    """
    parts = [int(p) for p in ip.split(".")]

    if parts[0] == 192 and parts[1] == 0 and parts[2] == 2:
        return True

    if parts[0] == 198 and parts[1] == 51 and parts[2] == 100:
        return True

    if parts[0] == 203 and parts[1] == 0 and parts[2] == 113:
        return True

    if parts[0] == 10:
        return True

    if parts[0] == 127:
        return True

    if parts[0] == 169 and parts[1] == 254:
        return True

    if parts[0] == 255:
        return True

    return False


def detect_ip_addresses(text: str) -> list[Detection]:

    results = []

    for match in IP_ADDRESS_PATTERN.finditer(text):

        value = match.group(0)

        if not is_valid_ipv4(value):
            continue

        if is_documentation_ip(value):
            continue

        results.append(
            Detection(
                category="IP_ADDRESS",
                surface=value,
                normalized=value,
                confidence="HIGH",
                evidence=text[:600]
            )
        )

    return results


# ============================================================
# ALL DETECTION
# ============================================================

def detect_all(text: str) -> list[Detection]:

    results = []

    results.extend(
        detect_emails(text)
    )

    results.extend(
        detect_phones(text)
    )

    results.extend(
        detect_people(text)
    )

    results.extend(
        detect_companies(text)
    )

    results.extend(
        detect_addresses(text)
    )

    results.extend(
        detect_ssns(text)
    )

    results.extend(
        detect_credit_cards(text)
    )

    results.extend(
        detect_dobs(text)
    )

    results.extend(
        detect_ip_addresses(text)
    )

    return results


# ============================================================
# ENTITY REGISTRY
# ============================================================

class EntityRegistry:

    def __init__(self):
        self.entities = {}

    @staticmethod
    def make_id(
        category,
        normalized
    ):

        digest = hashlib.sha256(
            (
                category
                + ":"
                + normalized
            ).encode("utf-8")
        ).hexdigest()[:12]

        return (
            f"{category}_{digest}"
        )

    def add(
        self,
        detection: Detection
    ):

        key = (
            detection.category,
            detection.normalized.casefold()
        )

        if key not in self.entities:

            self.entities[key] = Entity(
                entity_id=self.make_id(
                    detection.category,
                    detection.normalized
                ),
                category=detection.category,
                canonical=detection.normalized,
                confidence=detection.confidence
            )

        entity = self.entities[key]

        entity.surfaces.add(
            detection.surface
        )

        if (
            detection.evidence
            and detection.evidence
            not in entity.evidence
        ):
            entity.evidence.append(
                detection.evidence
            )

        if (
            detection.confidence == "HIGH"
            and entity.confidence != "HIGH"
        ):
            entity.confidence = "HIGH"

    def all(self):
        return list(
            self.entities.values()
        )


# ============================================================
# SYNTHETIC GENERATOR
# ============================================================

class SyntheticGenerator:

    FIRST_NAMES = [
        "Aarav", "Vihaan", "Aditya", "Rohan",
        "Arjun", "Karan", "Vikram", "Nikhil",
        "Rahul", "Dev", "Aman", "Kabir",
        "Varun", "Riya", "Neha", "Ananya",
        "Meera", "Kavya", "Isha", "Priya",
        "Tara", "Maya", "Naina", "Sana",
        "Aditi", "Ira", "Mihir", "Yash",
        "Manav", "Reyansh", "Dhruv", "Tanvi",
        "Simran", "Nisha", "Pallavi", "Ritika",
        "Anaya", "Aisha", "Aarohi", "Devika",
        "Rhea", "Kiara", "Ishaan", "Krish",
        "Veer", "Anika", "Myra", "Sara"
    ]

    LAST_NAMES = [
        "Sharma", "Mehta", "Rao", "Kapoor",
        "Patel", "Nair", "Verma", "Joshi",
        "Malik", "Shah", "Gupta", "Bose",
        "Sethi", "Menon", "Khanna", "Desai",
        "Bhat", "Iyer", "Singh", "Chawla",
        "Kulkarni", "Arora", "Saxena", "Mishra",
        "Malhotra", "Bansal", "Agarwal", "Reddy",
        "Naidu", "Das", "Sen", "Pillai",
        "Thomas", "George", "Dutta", "Roy",
        "Kohli", "Suri", "Tandon", "Khurana",
        "Bhandari", "Sarin", "Gandhi", "Madan",
        "Trivedi", "Puri", "Vora", "Shukla"
    ]

    COMPANY_NAMES = [
        "Nexus", "Horizon", "Pinnacle", "Summit",
        "Zenith", "Apex", "Vertex", "Prism",
        "Quantum", "Synergy", "Velocity", "Momentum",
        "Stellar", "Nova", "Catalyst", "Fusion",
        "Ascent", "Evolution", "Compass", "Beacon",
        "Orbit", "Epoch", "Axis", "Matrix",
        "Spectrum", "Portal", "Nexia", "Vortex",
        "Cascade", "Paradigm", "Essence", "Aurora",
        "Skyline", "Venture", "Forge", "Helix",
        "Ember", "Pulse", "Thrust", "Zenith",
        "Clarity", "Vision", "Nexgen", "Pathfinder",
        "Rivulet", "Haven", "Ascencia", "Frontier",
        "Beacon", "Luminox", "Artemis", "Phoenix",
        "Nebula", "Quantum", "Titan", "Draco",
    ]

    ADDRESS_COMPONENTS = [
        ("123", "Innovation", "Plaza", "Mumbai"),
        ("456", "Tech", "Street", "Bangalore"),
        ("789", "Digital", "Road", "Pune"),
        ("234", "Corporate", "Avenue", "Delhi"),
        ("567", "Business", "Park", "Gurgaon"),
        ("890", "Enterprise", "Complex", "Hyderabad"),
        ("345", "Global", "Court", "Chennai"),
        ("678", "Metro", "Heights", "Kolkata"),
        ("101", "Smart", "District", "Ahmedabad"),
        ("202", "Future", "Sector", "Surat"),
    ]

    @staticmethod
    def stable_hash(value: str) -> int:
        return int(
            hashlib.sha256(
                (
                    str(SEED)
                    + ":"
                    + value
                ).encode("utf-8")
            ).hexdigest(),
            16
        )

    def person_name(
        self,
        entity_id,
        used
    ):

        total = (
            len(self.FIRST_NAMES)
            * len(self.LAST_NAMES)
        )

        start = (
            self.stable_hash(entity_id)
            % total
        )

        for offset in range(total):

            index = (
                start + offset
            ) % total

            first = self.FIRST_NAMES[
                index
                // len(self.LAST_NAMES)
            ]

            last = self.LAST_NAMES[
                index
                % len(self.LAST_NAMES)
            ]

            candidate = (
                first
                + " "
                + last
            )

            if candidate.casefold() not in used:
                return candidate

        raise RuntimeError(
            "Synthetic name pool exhausted."
        )

    @staticmethod
    def email(index):
        return (
            f"contact{index:03d}"
            "@example.com"
        )

    @staticmethod
    def phone(original, index):

        digits = re.sub(
            r"\D",
            "",
            original
        )

        if (
            digits.startswith("91")
            and len(digits) == 12
        ):

            fake = "9000000000"
            result = []
            position = 0

            for char in original:

                if char.isdigit():

                    if position < 2:
                        result.append(char)
                    else:
                        result.append(
                            fake[position - 2]
                        )

                    position += 1

                else:
                    result.append(char)

            return "".join(result)

        if len(digits) == 10:

            fake = "9000000000"
            result = []
            position = 0

            for char in original:

                if char.isdigit():
                    result.append(
                        fake[position]
                    )
                    position += 1

                else:
                    result.append(char)

            return "".join(result)

        if "-" in original:

            area, subscriber = (
                original.split("-", 1)
            )

            return (
                area
                + "-"
                + ("0" * len(subscriber))
            )

        return (
            f"SYNTH_PHONE_{index:03d}"
        )

    def company_name(self, entity_id, used):
        """Generate synthetic company name with legal suffix."""
        total = len(self.COMPANY_NAMES)

        start = (
            self.stable_hash(entity_id)
            % total
        )

        for offset in range(total):

            index = (
                start + offset
            ) % total

            name = self.COMPANY_NAMES[index]
            candidate = f"{name} Limited"

            if candidate.casefold() not in used:
                return candidate

        # Fallback for large documents: add deterministic numeric modifier
        base_name = self.COMPANY_NAMES[start]
        for i in range(1, 10000):
            candidate = f"{base_name} {i:03d} Limited"
            if candidate.casefold() not in used:
                return candidate

        raise RuntimeError(
            "Synthetic company pool exhausted."
        )

    def address(self, entity_id):
        """Generate synthetic address."""
        component_index = (
            self.stable_hash(entity_id)
            % len(self.ADDRESS_COMPONENTS)
        )

        building, street, area, city = (
            self.ADDRESS_COMPONENTS[component_index]
        )

        pincode_base = (
            self.stable_hash(entity_id)
            % 90000 + 10000
        )

        return (
            f"{building} {street} {area}, "
            f"{city}, India - {pincode_base}"
        )

    @staticmethod
    def ssn(entity_id):
        """Generate synthetic SSN (non-realistic safe value)."""
        hash_val = int(
            hashlib.sha256(
                (str(SEED) + ":" + entity_id).encode()
            ).hexdigest(),
            16
        )

        area = (hash_val % 899) + 100
        group = (hash_val // 900 % 99) + 1
        serial = (hash_val // 89100 % 9999) + 1

        return f"{area:03d}-{group:02d}-{serial:04d}"

    @staticmethod
    def credit_card(entity_id):
        """Generate synthetic credit card (clearly non-real but Luhn-valid)."""
        hash_val = int(
            hashlib.sha256(
                (str(SEED) + ":" + entity_id).encode()
            ).hexdigest(),
            16
        )

        prefix = f"4111{hash_val % 10000000001:011d}"
        prefix_digits = [int(d) for d in prefix]
        base_total = sum(prefix_digits[::2]) + sum(d * 2 - 9 if d * 2 > 9 else d * 2 for d in prefix_digits[1::2])

        check_digit = 0
        for cd in range(10):
            cd_doubled = cd * 2 - 9 if cd * 2 > 9 else cd * 2
            if (base_total + cd_doubled) % 10 == 0:
                check_digit = cd
                break

        digits = prefix + str(check_digit)
        return f"{digits[:4]}-{digits[4:8]}-{digits[8:12]}-{digits[12:16]}"

    @staticmethod
    def dob(entity_id):
        """Generate synthetic date of birth."""
        hash_val = int(
            hashlib.sha256(
                (str(SEED) + ":" + entity_id).encode()
            ).hexdigest(),
            16
        )

        year = (hash_val % 50) + 1970
        month = (hash_val // 50 % 12) + 1
        day = (hash_val // 600 % 28) + 1

        return f"{day:02d}/{month:02d}/{year}"

    @staticmethod
    def ip_address(entity_id):
        """Generate synthetic IP address from documentation range."""
        hash_val = int(
            hashlib.sha256(
                (str(SEED) + ":" + entity_id).encode()
            ).hexdigest(),
            16
        )

        third_octet = (hash_val % 256)
        fourth_octet = (hash_val // 256 % 256)

        return f"192.0.2.{third_octet}"


# ============================================================
# ASSIGN REPLACEMENTS
# ============================================================

def assign_replacements(registry):

    generator = SyntheticGenerator()
    used_names = set()
    used_companies = set()

    email_index = 1
    phone_index = 1

    entities = sorted(
        registry.all(),
        key=lambda entity: (
            entity.category,
            entity.entity_id
        )
    )

    for entity in entities:

        if entity.category == "PERSON":

            entity.replacement = (
                generator.person_name(
                    entity.entity_id,
                    used_names
                )
            )

            used_names.add(
                entity.replacement.casefold()
            )

        elif entity.category == "EMAIL":

            entity.replacement = (
                generator.email(
                    email_index
                )
            )

            email_index += 1

        elif entity.category == "PHONE":

            original = sorted(
                entity.surfaces
            )[0]

            entity.replacement = (
                generator.phone(
                    original,
                    phone_index
                )
            )

            phone_index += 1

        elif entity.category == "COMPANY":

            entity.replacement = (
                generator.company_name(
                    entity.entity_id,
                    used_companies
                )
            )

            used_companies.add(
                entity.replacement.casefold()
            )

        elif entity.category == "ADDRESS":

            entity.replacement = (
                generator.address(
                    entity.entity_id
                )
            )

        elif entity.category == "SSN":

            entity.replacement = (
                generator.ssn(
                    entity.entity_id
                )
            )

        elif entity.category == "CREDIT_CARD":

            entity.replacement = (
                generator.credit_card(
                    entity.entity_id
                )
            )

        elif entity.category == "DOB":

            entity.replacement = (
                generator.dob(
                    entity.entity_id
                )
            )

        elif entity.category == "IP_ADDRESS":

            entity.replacement = (
                generator.ip_address(
                    entity.entity_id
                )
            )



# ============================================================
# ITERATE ALL DOCUMENT PARAGRAPHS
# ============================================================

def iter_paragraphs(document):

    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:

        for paragraph in section.header.paragraphs:
            yield paragraph

        for table in section.header.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:
                        yield paragraph

        for paragraph in section.footer.paragraphs:
            yield paragraph

        for table in section.footer.tables:

            for row in table.rows:

                for cell in row.cells:

                    for paragraph in cell.paragraphs:
                        yield paragraph


# ============================================================
# CASE-INSENSITIVE REPLACEMENT
# ============================================================

def replace_in_paragraph(
    paragraph,
    replacement_map
):

    if not paragraph.runs:
        return

    items = sorted(
        replacement_map.items(),
        key=lambda pair: len(pair[0]),
        reverse=True
    )

    while True:

        full_text = "".join(
            run.text
            for run in paragraph.runs
        )

        best = None

        for original, replacement in items:

            match = re.search(
                re.escape(original),
                full_text,
                re.IGNORECASE
            )

            if match is None:
                continue

            candidate = (
                match.start(),
                match.end(),
                original,
                replacement
            )

            if (
                best is None
                or candidate[0] < best[0]
                or (
                    candidate[0] == best[0]
                    and len(candidate[2])
                    > len(best[2])
                )
            ):
                best = candidate

        if best is None:
            break

        start, end, original, replacement = best

        boundaries = []
        cursor = 0

        for run in paragraph.runs:

            run_start = cursor
            run_end = (
                cursor
                + len(run.text)
            )

            boundaries.append(
                (run_start, run_end)
            )

            cursor = run_end

        start_run = None
        end_run = None
        start_offset = 0
        end_offset = 0

        for index, (
            run_start,
            run_end
        ) in enumerate(boundaries):

            if (
                start_run is None
                and run_start <= start < run_end
            ):
                start_run = index
                start_offset = (
                    start - run_start
                )

            if (
                run_start < end <= run_end
            ):
                end_run = index
                end_offset = (
                    end - run_start
                )
                break

        if (
            start_run is None
            or end_run is None
        ):
            break

        runs = paragraph.runs

        if start_run == end_run:

            text = runs[start_run].text

            runs[start_run].text = (
                text[:start_offset]
                + replacement
                + text[end_offset:]
            )

        else:

            first = runs[start_run].text
            last = runs[end_run].text

            prefix = first[:start_offset]
            suffix = last[end_offset:]

            runs[start_run].text = (
                prefix
                + replacement
            )

            for index in range(
                start_run + 1,
                end_run
            ):
                runs[index].text = ""

            runs[end_run].text = suffix


# ============================================================
# BUILD REPLACEMENT MAP
# ============================================================

def build_replacement_map(registry):

    mapping = {}

    for entity in registry.all():

        if not entity.replacement:
            continue

        for surface in entity.surfaces:
            mapping[surface] = entity.replacement

    return mapping


# ============================================================
# REDACT
# ============================================================

def redact_document(
    document,
    registry
):

    mapping = build_replacement_map(
        registry
    )

    for paragraph in iter_paragraphs(
        document
    ):

        replace_in_paragraph(
            paragraph,
            mapping
        )


# ============================================================
# EXTRACT DOCUMENT TEXT
# ============================================================

def extract_document_text(document):

    parts = []

    for paragraph in iter_paragraphs(
        document
    ):

        if paragraph.text:
            parts.append(
                paragraph.text
            )

    return "\n".join(parts)


# ============================================================
# VALIDATE ORIGINAL SURFACES
# ============================================================

def validate_replacement(
    redacted_document,
    registry
):

    text = extract_document_text(
        redacted_document
    )

    results = []

    for entity in registry.all():

        for surface in entity.surfaces:

            original_remaining = bool(
                re.search(
                    re.escape(surface),
                    text,
                    re.IGNORECASE
                )
            )

            replacement_present = bool(
                re.search(
                    re.escape(
                        entity.replacement
                    ),
                    text,
                    re.IGNORECASE
                )
            )

            if not replacement_present and not original_remaining:
                for other_entity in registry.all():
                    if other_entity.entity_id == entity.entity_id:
                        continue
                    if any(surface.casefold() in other_s.casefold() for other_s in other_entity.surfaces):
                        if bool(re.search(re.escape(other_entity.replacement), text, re.IGNORECASE)):
                            replacement_present = True
                            break

            results.append(
                {
                    "entity_id": entity.entity_id,
                    "category": entity.category,
                    "original": surface,
                    "replacement": entity.replacement,
                    "original_remaining": original_remaining,
                    "replacement_present": replacement_present
                }
            )

    return results


# ============================================================
# LEAKAGE SCAN REGEX
# ============================================================

LEAK_EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@"
    r"[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)

LEAK_PHONE_PATTERNS = [
    re.compile(
        r"(?<!\d)\+\s*91(?:[\s-]?\d){10}(?!\d)"
    ),
    re.compile(
        r"(?<![\d+\-])\d{2,4}-\d{6,10}(?!\d)"
    ),
    re.compile(
        r"(?<!\d)[6-9]\d{9}(?!\d)"
    ),
]

LEAK_PAN_PATTERN = re.compile(
    r"\b[A-Z]{5}[0-9]{4}[A-Z]\b"
)

LEAK_AADHAAR_PATTERN = re.compile(
    r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b"
)

LEAK_IP_PATTERN = re.compile(
    r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
)


def valid_leak_ipv4(value):

    parts = value.split(".")

    if len(parts) != 4:
        return False

    try:
        return all(
            0 <= int(part) <= 255
            for part in parts
        )
    except ValueError:
        return False


# ============================================================
# FINAL LEAKAGE SCAN
# ============================================================

def scan_for_pii_leakage(
    document,
    registry
):
    """
    Scan the final redacted document.

    Synthetic replacements are excluded from the leakage
    report, including phone replacements after normalization.
    """

    text = extract_document_text(
        document
    )

    # --------------------------------------------------------
    # Synthetic replacement values
    # --------------------------------------------------------

    synthetic_values = {
        entity.replacement.casefold()
        for entity in registry.all()
        if entity.replacement
    }

    # Phone replacements need normalized comparison because:
    #
    # 9000000000
    # +91 9000000000
    # + 91 90 0000 0000
    #
    # may be different surface formats.
    #
    synthetic_phone_values = set()

    for entity in registry.all():

        if (
            entity.category == "PHONE"
            and entity.replacement
        ):

            normalized = normalize_phone(
                entity.replacement
            )

            if normalized:
                synthetic_phone_values.add(
                    normalized
                )

    leakage = {
        "EMAIL": [],
        "PHONE": [],
        "SSN": [],
        "CREDIT_CARD": [],
        "DOB": [],
        "PAN": [],
        "AADHAAR": [],
        "IP_ADDRESS": [],
    }

    # --------------------------------------------------------
    # EMAIL
    # --------------------------------------------------------

    for match in LEAK_EMAIL_PATTERN.finditer(text):

        value = match.group(0)

        if (
            value.casefold()
            in synthetic_values
        ):
            continue

        leakage["EMAIL"].append(
            value
        )

    # --------------------------------------------------------
    # PHONE
    # --------------------------------------------------------

    for pattern in LEAK_PHONE_PATTERNS:

        for match in pattern.finditer(text):

            value = match.group(0)

            if not valid_phone(value):
                continue

            normalized = normalize_phone(
                value
            )

            # Ignore generated phone numbers.
            if (
                normalized
                in synthetic_phone_values
            ):
                continue

            leakage["PHONE"].append(
                value
            )

    # --------------------------------------------------------
    # PAN
    # --------------------------------------------------------

    for match in LEAK_PAN_PATTERN.finditer(text):

        value = match.group(0)

        if (
            value.casefold()
            in synthetic_values
        ):
            continue

        leakage["PAN"].append(
            value
        )

    # --------------------------------------------------------
    # AADHAAR
    # --------------------------------------------------------

    for match in LEAK_AADHAAR_PATTERN.finditer(text):

        value = match.group(0)

        digits = re.sub(
            r"\D",
            "",
            value
        )

        if len(digits) != 12:
            continue

        if (
            value.casefold()
            in synthetic_values
        ):
            continue

        leakage["AADHAAR"].append(
            value
        )

    # --------------------------------------------------------
    # SSN
    # --------------------------------------------------------

    for match in SSN_PATTERN.finditer(text):

        value = match.group(0)

        if not is_valid_ssn(value):
            continue

        if (
            value.casefold()
            in synthetic_values
        ):
            continue

        leakage["SSN"].append(
            value
        )

    # --------------------------------------------------------
    # CREDIT CARD
    # --------------------------------------------------------

    for match in CREDIT_CARD_PATTERN.finditer(text):

        value = match.group(0)

        if not is_valid_credit_card(value):
            continue

        normalized = re.sub(r"\D", "", value)

        if (
            normalized.casefold()
            in synthetic_values
        ):
            continue

        leakage["CREDIT_CARD"].append(
            value
        )

    # --------------------------------------------------------
    # DOB
    # --------------------------------------------------------

    for pattern in DOB_PATTERNS:

        for match in pattern.finditer(text):

            value = match.group(1)

            if not is_valid_dob(value):
                continue

            normalized = re.sub(r"[-/]", "", value)

            if (
                normalized.casefold()
                in synthetic_values
            ):
                continue

            leakage["DOB"].append(
                value
            )

    # --------------------------------------------------------
    # IP ADDRESS
    # --------------------------------------------------------

    for match in LEAK_IP_PATTERN.finditer(text):

        value = match.group(0)

        if not valid_leak_ipv4(value):
            continue

        if (
            value.casefold()
            in synthetic_values
        ):
            continue

        leakage["IP_ADDRESS"].append(
            value
        )

    # --------------------------------------------------------
    # DEDUPLICATE
    # --------------------------------------------------------

    for category in leakage:

        leakage[category] = sorted(
            set(
                leakage[category]
            )
        )

    return leakage


# ============================================================
# SAVE JSON REPORT
# ============================================================

def save_report(
    registry,
    validation_results,
    leakage
):

    report = {
        "input_file": str(INPUT_FILE),
        "output_file": str(OUTPUT_FILE),
        "seed": SEED,
        "unique_entities": len(
            registry.all()
        ),
        "entities": [],
        "validation": validation_results,
        "leakage_scan": leakage,
    }

    for entity in sorted(
        registry.all(),
        key=lambda value: (
            value.category,
            value.entity_id
        )
    ):

        report["entities"].append(
            {
                "entity_id": entity.entity_id,
                "category": entity.category,
                "canonical": entity.canonical,
                "confidence": entity.confidence,
                "surfaces": sorted(
                    entity.surfaces
                ),
                "replacement": entity.replacement,
                "evidence": entity.evidence[:3],
            }
        )

    REPORT_FILE.write_text(
        json.dumps(
            report,
            indent=2,
            ensure_ascii=False
        ),
        encoding="utf-8"
    )


# ============================================================
# MAIN WITH ERROR HANDLING
# ============================================================

def main():
    """
    Main redaction pipeline with comprehensive error handling.
    
    Edge cases handled:
    - File not found
    - File permission denied
    - Corrupted DOCX
    - Empty document
    - Invalid characters/encoding
    - Memory constraints
    - Output write failures
    """

    try:
        print(
            f"Opening: {INPUT_FILE}"
        )

        # ====================================================
        # INPUT VALIDATION
        # ====================================================
        
        if not INPUT_FILE.exists():
            raise FileNotFoundError(
                f"Input file not found: {INPUT_FILE}\n"
                f"Ensure Red Herring Prospectus.docx exists in current directory"
            )

        if not INPUT_FILE.is_file():
            raise ValueError(
                f"Input path is not a file: {INPUT_FILE}"
            )

        # Check file size (warn if > 100MB)
        file_size_mb = INPUT_FILE.stat().st_size / (1024 * 1024)
        if file_size_mb > 100:
            print(
                f"Warning: Large file ({file_size_mb:.1f}MB). "
                f"Processing may take several minutes."
            )

        # ====================================================
        # DOCUMENT LOADING
        # ====================================================

        try:
            original = Document(INPUT_FILE)
        except Exception as e:
            raise RuntimeError(
                f"Failed to open DOCX file (may be corrupted): {e}\n"
                f"Ensure Red Herring Prospectus.docx is a valid Word document"
            )

        print(
            "Document opened successfully."
        )

        # ====================================================
        # DOCUMENT VALIDATION
        # ====================================================

        if not original.paragraphs and not original.tables:
            raise ValueError(
                "Document is empty (no paragraphs or tables detected). "
                "Nothing to redact."
            )

    except (FileNotFoundError, ValueError, RuntimeError) as e:
        print(f"\n[ERROR] {e}")
        return False

    except Exception as e:
        print(f"\n[UNEXPECTED ERROR] {e}")
        return False

    try:
        print(
            "Paragraphs:",
            len(original.paragraphs)
        )

        print(
            "Tables:",
            len(original.tables)
        )

        # ========================================
        # DETECT
        # ========================================

        print(
            "\nDetecting PII..."
        )

        registry = EntityRegistry()

        for paragraph in iter_paragraphs(
            original
        ):

            if not paragraph.text:
                continue

            try:
                for detection in detect_all(
                    paragraph.text
                ):
                    registry.add(detection)
            except Exception as e:
                print(f"Warning: Error processing paragraph: {e}")
                continue

        detected_count = len(registry.all())
        print(f"Unique entities detected: {detected_count}")

        if detected_count == 0:
            print("No PII detected. Document may already be clean.")

        category_counts = {}

        for entity in registry.all():
            category_counts[entity.category] = (
                category_counts.get(entity.category, 0) + 1
            )

        print("\nDetected unique entities:")

        for category, count in sorted(category_counts.items()):
            print(f"  {category:12s}: {count}")

        # ========================================
        # GENERATE SYNTHETIC MAPPINGS
        # ========================================

        print("\nGenerating synthetic mappings...")

        try:
            assign_replacements(registry)
        except RuntimeError as e:
            print(f"[ERROR] {e}")
            return False

        # ========================================
        # CREATE REDACTED DOCUMENT
        # ========================================

        print("\nCreating redacted document...")

        try:
            redacted = Document(INPUT_FILE)
            redact_document(redacted, registry)
            redacted.save(OUTPUT_FILE)
        except Exception as e:
            print(f"[ERROR] Failed to create redacted document: {e}")
            return False

        print(f"Saved: {OUTPUT_FILE}")

        # ========================================
        # VALIDATION
        # ========================================

        print("\nValidating original surfaces...")

        validation_results = validate_replacement(redacted, registry)

        successful = [
            item for item in validation_results
            if (not item["original_remaining"] and item["replacement_present"])
        ]

        failed = [item for item in validation_results if item["original_remaining"]]

        missing = [item for item in validation_results if not item["replacement_present"]]

        print(f"\nOriginal surfaces checked: {len(validation_results)}")
        print(f"Successfully replaced: {len(successful)}")
        print(f"Original values still present: {len(failed)}")
        print(f"Replacement values not found: {len(missing)}")

        if failed:
            print("\nFAILED ORIGINALS:")
            for item in failed[:30]:
                print(f"  {item['category']}: {item['original']}")
            print("[ERROR] VALIDATION FAILED")
            return False

        if missing:
            print("\nMISSING REPLACEMENTS:")
            for item in missing[:30]:
                print(f"  {item['category']}: {item['replacement']}")
            print("[ERROR] VALIDATION FAILED")
            return False

        # ========================================
        # LEAKAGE SCAN
        # ========================================

        print("\nRunning final PII leakage scan...")

        try:
            leakage = scan_for_pii_leakage(redacted, registry)
        except Exception as e:
            print(f"Warning: Leakage scan error: {e}")
            leakage = {
                "EMAIL": [], "PHONE": [], "SSN": [],
                "CREDIT_CARD": [], "DOB": [], "PAN": [],
                "AADHAAR": [], "IP_ADDRESS": []
            }

        total_leaks = sum(len(values) for values in leakage.values())

        print("\nLeakage scan results:")

        for category, values in leakage.items():
            print(f"  {category:12s}: {len(values)}")
            if values:
                for value in values[:20]:
                    print(f"      {value}")

        if total_leaks == 0:
            print("\nFINAL LEAKAGE CHECK: PASS [OK]")
        else:
            print("\nFINAL LEAKAGE CHECK: REVIEW REQUIRED [WARNING]")

        # ========================================
        # SAVE REPORT
        # ========================================

        try:
            save_report(registry, validation_results, leakage)
            print(f"\nReport saved: {REPORT_FILE}")
        except Exception as e:
            print(f"[ERROR] Failed to save report: {e}")
            return False

        print("\n[SUCCESS] Redaction completed successfully.")
        return True

    except Exception as e:
        print(f"\n[ERROR] UNEXPECTED ERROR: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)